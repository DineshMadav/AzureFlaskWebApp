from markupsafe import escape
from werkzeug.utils import secure_filename
import docx
import time
from os.path import join
import sys, json

from flask import Flask, render_template, request, redirect, url_for, jsonify

# Functions

def savetest(test):
   	timestamp = "".join(((time.ctime()).replace(":","").split())[2:])        
	test_id = test['subject'] + "-" + timestamp
	#print('Saving Test ' + obj_test['subject'])

	path = "../../opentest"
	name = test_id + '.txt'  # Name of text file coerced with +.txt

	try:
		file = open(join(path, name),'w')   # Trying to create a new file or open one
		file.writelines(json.dumps(test))
		file.close()
		return test_id

	except:
		#print()
		#sys.exit(0) # quit Python
		return False




app = Flask(__name__)

@app.route("/")
def index():
    user = {'username': 'Miguel'}
    posts = [
        {
            'author': {'username': 'John'},
            'body': 'Beautiful day in Portland!'
        },
        {
            'author': {'username': 'Susan'},
            'body': 'The Avengers movie was so cool!'
        }
    ]
    return render_template('index.html', title='OpenTest', user=user, posts=posts)

@app.route('/user/<username>')
def show_user_profile(username):
    # show the user profile for that user
    return 'User %s' % escape(username)

@app.route('/post/<int:post_id>')
def show_post(post_id):
    # show the post with the given id, the id is an integer
    return 'Post %d' % post_id

@app.route('/path/<path:subpath>')
def show_subpath(subpath):
    # show the subpath after /path/
    return 'Subpath %s' % escape(subpath)

@app.route('/upload')
def upload_file():
   return render_template('upload.html')
"""	
@app.route('/uploader', methods = ['GET', 'POST'])
def save_file():
   if request.method == 'POST':
      f = request.files['file']
      #print(f.read())
      doc = docx.Document(f)
      #f.save(secure_filename(f.filename))
      objTest = []
      
      for para in doc.paragraphs:
         if len(para.text):
            objTest.append(para.text)
      #return 'file uploaded successfully'
      user = {'username': request.form['examiner']}
      #return render_template('index.html', title='OpenTest', user=user, posts=objTest)
      return redirect(url_for('confirm', examiner=user))
"""

@app.route('/uploader', methods = ['GET', 'POST'])
def save_file():
    if request.method == 'POST':
	req = request.form
	print(json.dumps(req.items()))
        f = request.files['file']
        #print(f.read())
        doc = docx.Document(f)
        
        #doc = docx.Document("testname.docx")
        objDoc = []

        for para in doc.paragraphs:
            if len(para.text):
                objDoc.append(para.text)

        #----------------------------------------------------------------------------

        testmeta = ("testname", "subject", "score", "time", "email")
        meta = dict(q = "questions", o = "options", a = "answers")
        objTest = {}

        objTest["questions"] = []
        objTest["options"] = []
        objTest["answers"] = []

        for line in objDoc:
            #print(line.split(":"))

            line_split = line.split(":")

            if (line_split[0]).strip() in ("q","o","a"):        
                objTest[meta[(line_split[0]).strip()]].append({"key":line_split[1].strip(),"value":line_split[2].strip()})
            else:
                objTest[(line_split[0]).strip()] = line_split[1].strip()

        #----------------------------------------------------------------------------

        total_q = len(objTest["questions"])

        objTestQuestions = {}

        #for item in test:
            #print(test[item])

        for i in range(total_q) :
            get_id = str(i + 1)

            for ques in objTest['questions']:
                if ques["key"] == get_id:
                    objTestQuestions[get_id] = {}
                    objTestQuestions[get_id]['question'] = ques["value"]

            options = []
            for opt in objTest['options']:
                if opt["key"] == get_id:
                    #print(opt["value"], end=", ")
                    options.append(opt["value"])
            objTestQuestions[get_id]['options'] = options

            for ans in objTest['answers']:
                if ans["key"] == get_id:
                    #objTestQuestions[get_id] = {}
                    objTestQuestions[get_id]['answer'] = ans["value"]

        objTest['all_questions_keys'] = objTestQuestions

        #----------------------------------------------------------------------------

        #total_q = len(objTest["questions"])

        objTestQuestionsList = []


        #for item in test:
            #print(test[item])

        for i in range(total_q) :
            get_id = str(i + 1)
            objTestQuestions = {}

            objTestQuestions['key'] = get_id
            #print(objTestQuestions)

            for ques in objTest['questions']:
                if ques["key"] == get_id:             
                    objTestQuestions['question'] = ques["value"]

            options = []
            for opt in objTest['options']:
                if opt["key"] == get_id:
                    #print(opt["value"], end=", ")
                    options.append(opt["value"])
            objTestQuestions['options'] = options

            for ans in objTest['answers']:
                if ans["key"] == get_id:
                    #objTestQuestions[get_id] = {}
                    objTestQuestions['answer'] = ans["value"]

            objTestQuestionsList.append(objTestQuestions)


        objTest['all_questions_display'] = objTestQuestionsList
	
	test_id = savetest(objTest)
	if test_id:
	   return render_template('confirmTest.html', title='OpenTest', objTest=objTest, test_id=test_id)
	else:
	   return "Cannot Save the Test. Please verify uploaded file (.doc or .docx) and try again"

"""     
@app.route('/savetest')
def callSaveTest():
	return savetest()
"""


@app.route('/confirm/<testid>/<action>')
def confirm(testid, action):
   if action == "Yes":
      #return "Test Paper saved successfully " + testid
      return redirect(url_for('viewtest', testid=testid))
   else:
      return "Test Paper deleted " + testid


@app.route('/viewtest/<testid>')
def viewtest(testid):
   try:
      objTest = json.loads((open("../../opentest/"+testid+".txt",'r')).read())
      return render_template('viewtest.html', title='OpenTest', objTest=objTest, test_id=testid)
   except:
      return "Cannot Open Test - "	+ testid	

@app.route('/starttest/<testid>')
def starttest(testid):
   try:
      objTest = json.loads((open("../../opentest/"+testid+".txt",'r')).read())
      return render_template('starttest.html', title='OpenTest', objTest=objTest, test_id=testid)
   except:
      return "Cannot Open Test - "	+ testid

@app.route('/submittest', methods = ['GET', 'POST'])
def submittest():
   if request.method == 'POST':
      req = request.form
      return jsonify(req.items())
   else:
      return "Invalid Submittion"
	




if __name__ == "__main__":
    app.run(debug = True)
